# Developed by Kevin Cai (kecai@cisco.com) and Sean Hwang (seahwang@cisco.com)
# Make sure [docx], [translate], [tkinter] are pip-installed.

import os
import docx
from translate import Translator
from tkinter import Tk, filedialog, messagebox

# OS User interface (displays the window)
root = Tk()         # Do not touch!
root.withdraw()     # DO not touch!

input_file_path = filedialog.askopenfilename(
    title="Select the input Word document",
    filetypes=[("Word Document", "*.docx")]
)

# File verification
# (file will only take .docx file in the first place, but this is another safety feature for stability)
if not input_file_path:
    print("Error: No input file selected.")
    exit(1)

if not input_file_path.endswith(".docx"):
    messagebox.showerror("Error", "Invalid file type. Please select a Word document (*.docx).")
    exit(1)

# Translation
input_doc = docx.Document(input_file_path)
translator = Translator(to_lang='zh')
translation = ''

for paragraph in input_doc.paragraphs:
    translation += translator.translate(paragraph.text) + '\n'

print("Translation:")
print(translation)

# Saving the translated file
save_translation = messagebox.askyesno("Save Translation", "Do you want to save the translated file?")

if save_translation:
    output_doc = docx.Document()
    output_doc.add_heading(input_file_path + " - Chinese Translation", level=1)
    output_doc.add_paragraph(translation)
    output_file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")]
    )
    if output_file_path:
        output_doc.save(output_file_path)
        print("Translation saved to:", output_file_path)
    else:
        print("Translation not saved.")
else:
    print("Translation:")
    print(translation)

print("Completed.")